from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from openpyxl import Workbook
from openpyxl import load_workbook 
from openpyxl.utils import column_index_from_string
from datetime import datetime 
import json
from datetime import datetime, timedelta
import tkinter as tk
from tkinter import filedialog 
import os
import time

# Data variables 
reporters = set()

class Job:
    def __init__(self, reporter, date, name, pages, rate, gross):
        self.reporter = reporter
        self.date = date
        self.name = name
        self.pages = pages
        self.rate = rate
        self.gross = gross

class Invoice:
    def __init__ (self, customer, jobs):
        self.customer = customer
        self.jobs = jobs

#runs on boot, initializing excel sheet
def Boot ():
    root = tk.Tk()
    root.withdraw()
    
    print ("Invoice Generator v1.1")
    print ("Select Excel File")



    filename = filedialog.askopenfilename(
        title="Select the Excel File",
        filetypes=[("Excel Files", "*.xlsx *.xlsm *.xltx *.xltm")]
    )

    wb = load_workbook(filename)
    ws = wb.active
    print (f"Loaded {filename} successfully!")
    reporterCol_letter = input("Reporter Column: ")
    dateCol_letter = input("Date Column: ")
    nameCol_letter = input("Name Column: ")
    pageNumCol_letter = input("Page Number Column: ")
    rateCol_letter = input("Rate Column: ")

    reporterCol = column_index_from_string(reporterCol_letter.upper())
    dateCol = column_index_from_string(dateCol_letter.upper())
    nameCol = column_index_from_string(nameCol_letter.upper())
    pageNumCol = column_index_from_string(pageNumCol_letter.upper())
    rateCol = column_index_from_string(rateCol_letter.upper())

    start = int(input("Start Row: "))
    end = int(input("End Row: "))
    return wb, ws, reporterCol, dateCol, nameCol, pageNumCol, rateCol, start, end
    
#finds all unique reporters
def findReporters (ws):
    for row in range (start, end + 1):
        cell_value = ws.cell(row, reporterCol).value
        if cell_value:
            reporters.add(cell_value)

#finds all jobs
def findJobs (ws):
    jobs = []
    
    for row in range (start, end + 1):
        reporter = ws.cell(row, reporterCol).value
        date = ws.cell(row, dateCol).value
        name = str(ws.cell(row, nameCol).value)
        pages = ws.cell(row, pageNumCol).value
        rate = ws.cell(row, rateCol).value
        gross = pages * rate

        job = Job(reporter, date, name, pages, rate, gross)

        jobs.append(job)
    
    return jobs

#groups all jobs by reporter
def groupJobs ():
    reporter_jobs = {}
    for job in jobs:
        if job.reporter not in reporter_jobs:
            reporter_jobs[job.reporter] = []
        reporter_jobs[job.reporter].append(job)
    return reporter_jobs

#create data for invoices   
def createInvoices():
    invoices = []
    for reporter in reporter_jobs:
        invoice = Invoice(reporter, [])
        invoice.customer = reporter
        for job in reporter_jobs[reporter]:
            invoice.jobs.append(job)
        invoices.append(invoice)
    return invoices

#publish invoices
def publishInvoices():
    root = tk.Tk()
    root.withdraw()
    print ("Select DOCX Folder")
    folderName = filedialog.askdirectory(title="Select DOCX Folder")

    startNum = int(input("Enter Starting Number: "))
    currentIndex = startNum
    grandTotal = 0

    for invoice in invoices:
        
        #initialize the document + header
        doc = Document()
        with open("header.json") as f:
            header_data = json.load(f)
        today = datetime.today()
        two_weeks = today + timedelta(days=14)
        headingText = f"Invoice #{currentIndex}\n\n\n\n{header_data['Company name']}\n{header_data['Address1']}\n{header_data['Address2']}\n{header_data['Cell']}\n\n" \
            f"TO: {invoice.customer}\n\n\n\n\n\nInvoice Date: {today.strftime('%m-%d-%Y')}\nDate Due: {two_weeks.strftime('%m-%d-%Y')}\n\n"
        heading = doc.add_table(rows=1, cols=5)
        heading.columns[1].width = Inches(1.8)
        heading.columns[0].width = Inches(2.0)
        heading.alignment = WD_TABLE_ALIGNMENT.CENTER
        heading.rows[0].cells[0].text = headingText


        
        #intialize table
        table = doc.add_table(rows=1, cols=5)
        table.columns[1].width = Inches(2.5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_cells = table.rows[0].cells
        header_cells[0].text = "Date"
        header_cells[1].text = "Job Title"
        header_cells[2].text = "Pages"
        header_cells[3].text = "Rate"
        header_cells[4].text = "Total"
        for cell in table.rows[0].cells:
            cell.text = cell.text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        #add jobs dynamically to the table
        for job in invoice.jobs:
            row_cells = table.add_row().cells
            row_cells[0].text = job.date.strftime("%m-%d-%Y")
            row_cells[1].text = job.name
            row_cells[2].text = str(job.pages)
            row_cells[3].text = f"${job.rate:.2f}/pg"
            row_cells[4].text = f"${job.gross:.2f}"

            

        #add total and footer
        invoice_total = sum(job.pages * job.rate for job in invoice.jobs)
        total = doc.add_paragraph(f"Total Due:  ${invoice_total:.2f}", style='Normal')
        total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in total.runs:
            run.bold = True

        footer = f"{header_data['Footer1']}\n{header_data['Footer2']}"
        p2 = doc.add_paragraph(footer, style='Normal')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        #determine file name and export, moving onto the next job
        reporterInitials = ''.join([part[0].upper() for part in invoice.customer.split()])
        fileName = os.path.join(folderName, f"Invoice{currentIndex}-{reporterInitials}.docx")
        doc.save(fileName)
        print(f"Succesfully Generated {invoice.customer}'s Invoice! - ${invoice_total:.2f}")
        currentIndex += 1
        grandTotal += invoice_total
    return grandTotal, folderName

def createPDFs(folderName):
    root = tk.Tk()
    root.withdraw()
    print ("Select PDF Folder")
    pdfFolderName = filedialog.askdirectory(title="Select PDF Folder")

    for file in os.listdir(folderName):
        if file.endswith(".docx"):
            src = os.path.join(folderName, file)
            dst = os.path.join(pdfFolderName, file.replace(".docx", ".pdf"))
            try:
                convert(src, dst)
                print(f"Converted: {file}")
                time.sleep(0.3)  # prevents Word from stalling
            except Exception as e:
                print(f"Failed: {file} — {e}")
    print("Succesfully Converted to PDFS")




wb, ws, reporterCol, dateCol, nameCol, pageNumCol, rateCol, start, end = Boot()
findReporters(ws)
jobs = findJobs(ws)
reporter_jobs = groupJobs()
invoices = createInvoices()
grandTotal, folderName = publishInvoices()
createPDFs(folderName)
print(f"Grand Total: ${grandTotal}")
input("Press Enter to exit")

#FIXME: header is still weird, only create pdfs of files created in an instance of the program